#!/usr/bin/env python3
"""
Ingestão da Biblioteca -> ChromaDB
Lê recursivamente as pastas configuradas, extrai texto, faz chunking,
gera embeddings via Ollama (nomic-embed-text) e indexa no vetorial local.

Paralelismo:
    - Extração de texto roda em processos separados (ProcessPoolExecutor),
      usando vários núcleos de CPU ao mesmo tempo em vez de 1 arquivo por vez.
    - Vários arquivos são "indexados" (chunk + embedding + gravação no banco)
      ao mesmo tempo, cada um com sua própria barra de progresso.
    - A geração de embeddings em si roda em threads (ThreadPoolExecutor)
      compartilhadas entre todos os arquivos ativos, disparando várias
      chamadas simultâneas ao Ollama para aproveitar a folga da GPU.

Uso:
    python ingest.py                          # ingere tudo (full rescan)
    python ingest.py --watch                  # fica rodando e reingere o que mudar
    python ingest.py --extract-workers 4 --file-workers 3 --embed-workers 8
"""

import os
import sys
import time
import logging
import hashlib
import argparse
import threading
from pathlib import Path
from datetime import datetime
from concurrent.futures import (
    ThreadPoolExecutor,
    ProcessPoolExecutor,
    TimeoutError as FutureTimeoutError,
    as_completed,
    wait as futures_wait,
    FIRST_COMPLETED,
)

import chromadb
import ollama
from tqdm import tqdm
from pypdf import PdfReader
from docx import Document as DocxDocument
from odf import text as odf_text, teletype
from odf.opendocument import load as odf_load
import ebooklib
from ebooklib import epub
from bs4 import BeautifulSoup

# O pypdf usa o módulo "logging" pra avisar sobre PDFs mal formados
# ("invalid pdf header", "incorrect startxref pointer", etc). Isso é só
# ruído pra gente (o pypdf tenta se recuperar sozinho) e bagunça as barras
# de progresso no terminal — então silenciamos, deixando só erros graves.
logging.getLogger("pypdf").setLevel(logging.ERROR)

# ---------------- CONFIG ----------------
YURIAI = Path("C:/IA")
VECTOR_DB_PATH = YURIAI / "AI" / "environment" / "vector-db"
COLLECTION_NAME = "biblioteca"      # coleção única — a fonte fica no metadado "source_type"
EMBED_MODEL = "nomic-embed-text"   # rode: ollama pull nomic-embed-text
CHUNK_SIZE = 1000       # caracteres por chunk
CHUNK_OVERLAP = 150
SUPPORTED_EXT = {".txt", ".md", ".pdf", ".docx", ".odt", ".epub"}
CHUNK_EMBED_TIMEOUT_SECONDS = 45     # timeout por CHAMADA individual de embedding
EXTRACTION_TIMEOUT_SECONDS = 90      # timeout da extração de um arquivo (proteção contra PDF corrompido)
SKIPPED_LOG_PATH = YURIAI / "AI" / "environment" / "rag" / "skipped_files.log"

# --------------- PARALELISMO ---------------
# Extração de texto é CPU-bound (pypdf/docx/odf) -> processos, um por núcleo.
EXTRACT_WORKERS = os.cpu_count() or 4
# Quantos arquivos ficam sendo "indexados" (chunk + embedding + gravação)
# ao mesmo tempo -> cada um ganha sua própria barra de progresso na tela.
FILE_WORKERS = 3
# Chamadas de embedding simultâneas ao Ollama, COMPARTILHADAS entre todos os
# arquivos ativos ao mesmo tempo (é o que de fato ocupa a GPU). Ajuste
# conforme o quanto a GPU aguenta antes de saturar (veja a nota sobre
# OLLAMA_NUM_PARALLEL no final do script).
EMBED_WORKERS = 8
# -----------------------------------------


def log_skipped(path: Path, reason: str):
    """Registra no console E num arquivo de log persistente."""
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    line = f"[{timestamp}] PULADO: {path} — motivo: {reason}"
    tqdm.write(f"  ⚠️  {line}")
    try:
        SKIPPED_LOG_PATH.parent.mkdir(parents=True, exist_ok=True)
        with open(SKIPPED_LOG_PATH, "a", encoding="utf-8") as f:
            f.write(line + "\n")
    except Exception:
        pass  # não deixa falha no log derrubar a indexação


def extract_text(path: Path) -> str:
    """Extrai texto puro de um arquivo suportado.
    Roda dentro de um worker do ProcessPoolExecutor — precisa continuar
    sendo uma função de nível de módulo (picklable) e não depender de
    nenhum estado global (chromadb, ollama, etc.)."""
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        try:
            reader = PdfReader(str(path))
            if reader.is_encrypted:
                # Muitos PDFs de apostila/cursinho vêm com criptografia só pra
                # bloquear impressão/cópia, sem senha real pra abrir — uma
                # senha vazia costuma destravar a leitura nesses casos.
                try:
                    reader.decrypt("")
                except Exception:
                    pass
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception as e:
            print(f"  [erro ao ler PDF] {path.name}: {e}")
            return ""

    elif suffix == ".docx":
        try:
            doc = DocxDocument(str(path))
            partes = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    partes.append(" | ".join(cell.text for cell in row.cells))
            return "\n".join(partes)
        except Exception as e:
            print(f"  [erro ao ler DOCX] {path.name}: {e}")
            return ""

    elif suffix == ".odt":
        try:
            doc = odf_load(str(path))
            paragraphs = doc.getElementsByType(odf_text.P)
            return "\n".join(teletype.extractText(p) for p in paragraphs)
        except Exception as e:
            print(f"  [erro ao ler ODT] {path.name}: {e}")
            return ""

    elif suffix == ".epub":
        try:
            book = epub.read_epub(str(path), options={"ignore_ncx": True})
            partes = []
            for item in book.get_items():
                if item.get_type() == ebooklib.ITEM_DOCUMENT:
                    soup = BeautifulSoup(item.get_content(), "html.parser")
                    texto = soup.get_text(separator="\n")
                    if texto.strip():
                        partes.append(texto)
            return "\n\n".join(partes)
        except Exception as e:
            print(f"  [erro ao ler EPUB] {path.name}: {e}")
            return ""

    else:
        try:
            return path.read_text(encoding="utf-8", errors="ignore")
        except Exception as e:
            print(f"  [erro ao ler] {path.name}: {e}")
            return ""


def chunk_text(text: str, size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP):
    """Chunking simples por caracteres com overlap."""
    text = text.strip()
    if not text:
        return []
    chunks = []
    start = 0
    while start < len(text):
        end = start + size
        chunks.append(text[start:end])
        start = end - overlap
    return chunks


def file_hash(path: Path) -> str:
    """Hash do conteúdo do arquivo, pra saber se mudou desde a última ingestão."""
    h = hashlib.sha256()
    h.update(path.read_bytes())
    return h.hexdigest()[:16]


def get_collection():
    client = chromadb.PersistentClient(path=str(VECTOR_DB_PATH))
    return client.get_or_create_collection(name=COLLECTION_NAME)


# Formatos imutáveis: se já foi indexado uma vez, nunca precisa reconferir conteúdo.
IMMUTABLE_EXT = {".pdf", ".epub"}
# Formatos editáveis: sempre reconferir hash pra detectar mudanças de conteúdo.
EDITABLE_EXT = {".docx", ".odt", ".doc", ".txt", ".md"}


def already_indexed_by_id(collection, doc_id: str) -> bool:
    """Verifica só se o doc_id já existe no índice, sem olhar conteúdo (usado para PDFs)."""
    try:
        existing = collection.get(where={"doc_id": doc_id}, limit=1)
        return bool(existing["ids"])
    except Exception:
        return False


def already_indexed(collection, doc_id: str, current_hash: str) -> bool:
    """Verifica se esse arquivo (com esse hash) já está indexado."""
    try:
        existing = collection.get(where={"doc_id": doc_id}, limit=1)
        if existing["ids"]:
            return existing["metadatas"][0].get("file_hash") == current_hash
    except Exception:
        pass
    return False


def remove_old_chunks(collection, doc_id: str):
    try:
        collection.delete(where={"doc_id": doc_id})
    except Exception:
        pass


def determine_ingest_targets(files, root: Path, source_tag: str, collection):
    """Filtra a lista de arquivos, mantendo só os que precisam ser (re)indexados.
    Isso é só metadado no Chroma + hash de arquivo, então roda sequencial
    mesmo (é rápido por arquivo, mas soma tempo com milhares de arquivos) —
    a parte pesada (extração + embedding) é paralelizada depois.
    Mostra uma barra própria pra essa fase, marcando cada arquivo já
    indexado como "concluído" instantaneamente, pra não parecer que nada
    está acontecendo enquanto os milhares de arquivos são conferidos.
    Retorna lista de tuplas (path, doc_id, current_hash)."""
    targets = []
    already = 0
    novos = 0

    check_bar = tqdm(files, desc="Conferindo já indexados", unit="arquivo", dynamic_ncols=True)
    for path in check_bar:
        doc_id = f"{source_tag}::{path.relative_to(root)}"
        suffix = path.suffix.lower()

        if suffix in IMMUTABLE_EXT:
            if already_indexed_by_id(collection, doc_id):
                already += 1
                check_bar.set_postfix(já_indexados=already, novos=novos)
                continue
            targets.append((path, doc_id, None))
        else:
            try:
                current_hash = file_hash(path)
            except (FileNotFoundError, OSError, PermissionError) as e:
                print(f"  [pulado — não foi possível ler] {path.relative_to(root)}: {e}")
                continue
            if already_indexed(collection, doc_id, current_hash):
                already += 1
                check_bar.set_postfix(já_indexados=already, novos=novos)
                continue
            targets.append((path, doc_id, current_hash))

        novos += 1
        check_bar.set_postfix(já_indexados=already, novos=novos)

    check_bar.close()
    print(f"Checagem concluída: {already} já indexado(s) (pulados), {novos} novo(s)/alterado(s) a processar.")
    return targets, already


class PositionPool:
    """Distribui as linhas da tela (posições do tqdm) entre os arquivos que
    estão sendo indexados ao mesmo tempo, pra cada um ter sua barra fixa
    numa linha só (em vez de ficarem se sobrepondo/pulando)."""

    def __init__(self, size: int, start: int = 1):
        self._lock = threading.Lock()
        self._free = list(range(start, start + size))  # posições abaixo de `start` são reservadas pras barras fixas

    def acquire(self):
        with self._lock:
            return self._free.pop(0) if self._free else None

    def release(self, pos):
        if pos is None:
            return
        with self._lock:
            self._free.append(pos)
            self._free.sort()


def embed_chunks_parallel(path: Path, chunks, embed_executor: ThreadPoolExecutor, progress_desc: str,
                           position=None, max_in_flight=None):
    """Gera os embeddings de todos os chunks de um arquivo usando o pool de
    threads COMPARTILHADO entre todos os arquivos ativos.

    Importante: em vez de jogar todos os chunks na fila de uma vez (o que
    fazia um arquivo com milhares de chunks entupir a fila e travar de
    verdade os arquivos pequenos que chegassem depois — eles ficavam presos
    atrás de uma fila gigante e estouravam o timeout mesmo sem o Ollama
    estar lento), mantemos só uma "janela" de no máximo `max_in_flight`
    pedidos em voo por arquivo. Assim que um termina, o próximo da fila
    desse arquivo é liberado. Isso deixa vários arquivos dividirem a fila
    de forma justa, mesmo que um deles tenha muito mais chunks que os outros.

    O timeout aqui detecta ESTAGNAÇÃO: se em nenhum momento, por
    CHUNK_EMBED_TIMEOUT_SECONDS segundos, nenhum dos pedidos em voo termina,
    assumimos que o Ollama travou de verdade e pulamos o arquivo inteiro
    (nunca indexamos parcialmente)."""
    n = len(chunks)
    if max_in_flight is None:
        max_in_flight = embed_executor._max_workers
    max_in_flight = max(1, min(max_in_flight, n))

    embeddings = [None] * n
    next_idx = 0
    in_flight = {}  # future -> índice do chunk

    def submit_next():
        nonlocal next_idx
        if next_idx < n:
            fut = embed_executor.submit(ollama.embeddings, model=EMBED_MODEL, prompt=chunks[next_idx])
            in_flight[fut] = next_idx
            next_idx += 1

    chunk_bar = tqdm(
        total=n, desc=f"    {progress_desc[:40]}", unit="chunk",
        leave=False, position=position, dynamic_ncols=True,
    )
    try:
        for _ in range(max_in_flight):
            submit_next()

        while in_flight:
            done, _ = futures_wait(
                list(in_flight.keys()), timeout=CHUNK_EMBED_TIMEOUT_SECONDS, return_when=FIRST_COMPLETED
            )
            if not done:
                stuck_idx = min(in_flight.values())
                log_skipped(
                    path,
                    f"embedding travou no chunk {stuck_idx+1}/{n} "
                    f"(>{CHUNK_EMBED_TIMEOUT_SECONDS}s sem NENHUM progresso) — arquivo pulado inteiro"
                )
                for f in in_flight:
                    f.cancel()
                return None

            for fut in done:
                idx = in_flight.pop(fut)
                try:
                    result = fut.result()
                    embeddings[idx] = result["embedding"]
                    chunk_bar.update(1)
                except Exception as e:
                    log_skipped(path, f"erro ao gerar embedding do chunk {idx+1}/{n} — {e} — arquivo pulado inteiro")
                    for f in in_flight:
                        f.cancel()
                    return None
                submit_next()
    finally:
        chunk_bar.close()

    return embeddings


def process_one_file(path, doc_id, current_hash, text, source_tag, collection,
                      embed_pool, position_pool, root, write_lock, counters, counters_lock,
                      session_bar, biblioteca_bar):
    """Faz o chunk + embedding + gravação de UM arquivo. Roda dentro do
    file_pool, então vários arquivos passam por aqui ao mesmo tempo.

    session_bar: progresso desta sessão (só os arquivos novos/alterados, 0->100%).
    biblioteca_bar: progresso da Biblioteca inteira (já parte de "já_indexados"
    pré-preenchido, e só avança quando um arquivo é indexado com SUCESSO —
    arquivos pulados por erro não contam como "completos")."""
    chunks = chunk_text(text)
    if not chunks:
        with counters_lock:
            session_bar.update(1)
        return

    pos = position_pool.acquire()
    tqdm.write(f"  Indexando [{source_tag}]: {path.relative_to(root)} ({len(chunks)} chunks)")
    try:
        embeddings = embed_chunks_parallel(path, chunks, embed_pool, path.name, position=pos)
    finally:
        position_pool.release(pos)

    if embeddings is None:
        with counters_lock:
            counters["skipped"] += 1
            session_bar.update(1)
        return

    ids = [f"{doc_id}::{i}" for i in range(len(chunks))]
    metadatas = [{
        "doc_id": doc_id,
        "source_type": source_tag,       # "biblioteca" ou "affine"
        "file_hash": current_hash or "imutavel",
        "source_path": str(path),
        "chunk_index": i,
    } for i in range(len(chunks))]

    # Escritas no Chroma ficam serializadas (é rápido, o gargalo é o
    # embedding, não a gravação) pra evitar concorrência no banco.
    with write_lock:
        remove_old_chunks(collection, doc_id)
        collection.add(ids=ids, embeddings=embeddings, documents=chunks, metadatas=metadatas)

    with counters_lock:
        counters["updated"] += 1
        session_bar.update(1)
        biblioteca_bar.update(1)  # só avança na barra geral quando realmente indexado


def scan_and_ingest(root: Path, source_tag: str,
                     extract_pool: ProcessPoolExecutor, embed_pool: ThreadPoolExecutor, file_pool: ThreadPoolExecutor):
    if not root.exists():
        print(f"Pasta não encontrada: {root}")
        sys.exit(1)

    collection = get_collection()
    files = [p for p in root.rglob("*") if p.is_file() and p.suffix.lower() in SUPPORTED_EXT]

    if not files:
        print(f"Nenhum arquivo suportado encontrado em {root} (tipos: {SUPPORTED_EXT})")
        return

    print(f"[{source_tag}] Encontrados {len(files)} arquivos. Verificando mudanças...")
    targets, already = determine_ingest_targets(files, root, source_tag, collection)

    if not targets:
        print(f"Nada novo para indexar [{source_tag}]. Total de chunks na coleção: {collection.count()}")
        return

    n_extract = extract_pool._max_workers
    n_files = file_pool._max_workers
    n_embed = embed_pool._max_workers
    print(
        f"[{source_tag}] {len(targets)} arquivo(s) para (re)indexar — "
        f"até {n_extract} extraindo texto em paralelo, "
        f"até {n_files} arquivo(s) sendo indexado(s) ao mesmo tempo, "
        f"{n_embed} chamada(s) de embedding simultâneas no total."
    )

    counters = {"updated": 0, "skipped": 0}
    counters_lock = threading.Lock()
    write_lock = threading.Lock()
    position_pool = PositionPool(n_files, start=2)  # posições 0 e 1 são as duas barras fixas

    # Barra 0: Biblioteca inteira — já nasce preenchida com o que estava indexado antes.
    # Só avança quando um arquivo é indexado com SUCESSO (não conta pulados como "completo").
    biblioteca_bar = tqdm(total=len(files), initial=already, desc="Biblioteca completa", unit="arquivo",
                           position=0, leave=True, dynamic_ncols=True)
    # Barra 1: progresso desta sessão — só os arquivos novos/alterados, sempre 0% -> 100%.
    session_bar = tqdm(total=len(targets), desc="Progresso da sessão", unit="arquivo",
                        position=1, leave=True, dynamic_ncols=True)

    # Dispara a extração de TODOS os arquivos de uma vez no pool de processos.
    future_to_target = {
        extract_pool.submit(extract_text, path): (path, doc_id, current_hash)
        for path, doc_id, current_hash in targets
    }

    file_tasks = []
    for future in as_completed(future_to_target):
        path, doc_id, current_hash = future_to_target[future]

        try:
            text = future.result(timeout=EXTRACTION_TIMEOUT_SECONDS)
        except FutureTimeoutError:
            log_skipped(path, f"extração de texto travou (>{EXTRACTION_TIMEOUT_SECONDS}s) — provável arquivo corrompido")
            with counters_lock:
                counters["skipped"] += 1
                session_bar.update(1)  # só a sessão avança; biblioteca_bar não, pois não foi indexado de fato
            continue
        except Exception as e:
            log_skipped(path, f"erro inesperado na extração — {e}")
            with counters_lock:
                counters["skipped"] += 1
                session_bar.update(1)
            continue

        # Cada arquivo pronto (texto extraído) vai pro pool de indexação —
        # até n_files arquivos processados (chunk + embedding + gravação)
        # ao mesmo tempo, cada um com sua barra própria.
        task = file_pool.submit(
            process_one_file, path, doc_id, current_hash, text, source_tag, collection,
            embed_pool, position_pool, root, write_lock, counters, counters_lock,
            session_bar, biblioteca_bar,
        )
        file_tasks.append(task)

    for task in file_tasks:
        task.result()  # espera terminar e propaga qualquer exceção inesperada

    session_bar.close()
    biblioteca_bar.close()

    print(f"Concluído. {counters['updated']} arquivo(s) atualizado(s)/adicionado(s), "
          f"{counters['skipped']} pulado(s) por erro [{source_tag}].")
    print(f"Total de chunks na coleção (todas as fontes): {collection.count()}")
    if counters["skipped"]:
        print(f"Detalhes dos pulados salvos em: {SKIPPED_LOG_PATH}")


def watch_mode(root: Path, source_tag: str, extract_pool, embed_pool, file_pool, interval=60):
    print(f"Modo watch ativo em [{source_tag}] — reingestão a cada {interval}s. Ctrl+C pra parar.")
    while True:
        scan_and_ingest(root, source_tag, extract_pool, embed_pool, file_pool)
        time.sleep(interval)


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--root", type=str, default=str(YURIAI / "Biblioteca"),
                         help="Pasta raiz a indexar (padrão: Biblioteca)")
    parser.add_argument("--source-tag", type=str, default="biblioteca",
                         help="Rótulo da fonte, ex: biblioteca, affine")
    parser.add_argument("--watch", action="store_true", help="Fica rodando e reingere periodicamente")
    parser.add_argument("--interval", type=int, default=60, help="Intervalo do watch em segundos")
    parser.add_argument("--extract-workers", type=int, default=EXTRACT_WORKERS,
                         help=f"Processos paralelos para extração de texto (padrão: {EXTRACT_WORKERS}, = núcleos de CPU)")
    parser.add_argument("--file-workers", type=int, default=FILE_WORKERS,
                         help=f"Quantos arquivos são indexados ao mesmo tempo (padrão: {FILE_WORKERS})")
    parser.add_argument("--embed-workers", type=int, default=EMBED_WORKERS,
                         help=f"Chamadas de embedding simultâneas ao Ollama, no total (padrão: {EMBED_WORKERS})")
    args = parser.parse_args()

    root_path = Path(args.root)

    # Pools persistentes, criados uma vez e reaproveitados em todas as
    # iterações (inclusive no --watch).
    with ProcessPoolExecutor(max_workers=args.extract_workers) as extract_pool, \
         ThreadPoolExecutor(max_workers=args.embed_workers) as embed_pool, \
         ThreadPoolExecutor(max_workers=args.file_workers) as file_pool:
        try:
            if args.watch:
                watch_mode(root_path, args.source_tag, extract_pool, embed_pool, file_pool, args.interval)
            else:
                scan_and_ingest(root_path, args.source_tag, extract_pool, embed_pool, file_pool)
        except KeyboardInterrupt:
            print("\nInterrompido pelo usuário.")
