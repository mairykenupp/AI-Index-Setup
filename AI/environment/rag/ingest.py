#!/usr/bin/env python3
"""
Ingestão da Biblioteca -> ChromaDB
Lê recursivamente as pastas configuradas, extrai texto, faz chunking,
gera embeddings via Ollama (nomic-embed-text) e indexa no vetorial local.

Uso:
    python ingest.py                # ingere tudo (full rescan)
    python ingest.py --watch        # fica rodando e reingere o que mudar
"""

import os
import sys
import time
import hashlib
import argparse
from pathlib import Path
from datetime import datetime
from concurrent.futures import ThreadPoolExecutor, TimeoutError as FutureTimeoutError

import chromadb
import ollama
from tqdm import tqdm
from pypdf import PdfReader
from docx import Document as DocxDocument
from odf import text as odf_text, teletype
from odf.opendocument import load as odf_load

# ---------------- CONFIG ----------------
YURIAI = Path("C:/IA")
VECTOR_DB_PATH = YURIAI / "AI" / "environment" / "vector-db"
COLLECTION_NAME = "biblioteca"      # coleção única — a fonte fica no metadado "source_type"
EMBED_MODEL = "nomic-embed-text"   # rode: ollama pull nomic-embed-text
CHUNK_SIZE = 1000       # caracteres por chunk
CHUNK_OVERLAP = 150
SUPPORTED_EXT = {".txt", ".md", ".pdf", ".docx", ".odt"}
CHUNK_EMBED_TIMEOUT_SECONDS = 45     # timeout por CHAMADA individual de embedding (não pelo arquivo inteiro)
EXTRACTION_TIMEOUT_SECONDS = 90      # timeout só pra etapa de extrair texto do arquivo (proteção contra PDF corrompido travando o pypdf)
SKIPPED_LOG_PATH = YURIAI / "AI" / "environment" / "rag" / "skipped_files.log"
# -----------------------------------------


def run_with_timeout(func, timeout, *args, **kwargs):
    """Roda func com um timeout real, sem depender de progresso acumulado.
    Levanta FutureTimeoutError se a chamada específica não responder a tempo.
    A thread travada (se houver) fica órfã em segundo plano e não bloqueia o resto."""
    executor = ThreadPoolExecutor(max_workers=1)
    try:
        future = executor.submit(func, *args, **kwargs)
        return future.result(timeout=timeout)
    finally:
        executor.shutdown(wait=False)


def log_skipped(path: Path, reason: str):
    """Registra no console E num arquivo de log persistente."""
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    line = f"[{timestamp}] PULADO: {path} — motivo: {reason}"
    tqdm.write(f"  ⚠️  {line}")
    try:
        SKIPPED_LOG_PATH.parent.mkdir(parents=True, exist_ok=True)
        with open(SKIPPED_LOG_PATH, "a", encoding="utf-8") as f:
            f.write(line + "\n")
    except Exception:
        pass  # não deixa falha no log derrubar a indexação


def extract_text(path: Path) -> str:
    """Extrai texto puro de um arquivo suportado."""
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        try:
            reader = PdfReader(str(path))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception as e:
            print(f"  [erro ao ler PDF] {path.name}: {e}")
            return ""

    elif suffix == ".docx":
        try:
            doc = DocxDocument(str(path))
            partes = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    partes.append(" | ".join(cell.text for cell in row.cells))
            return "\n".join(partes)
        except Exception as e:
            print(f"  [erro ao ler DOCX] {path.name}: {e}")
            return ""

    elif suffix == ".odt":
        try:
            doc = odf_load(str(path))
            paragraphs = doc.getElementsByType(odf_text.P)
            return "\n".join(teletype.extractText(p) for p in paragraphs)
        except Exception as e:
            print(f"  [erro ao ler ODT] {path.name}: {e}")
            return ""

    else:
        try:
            return path.read_text(encoding="utf-8", errors="ignore")
        except Exception as e:
            print(f"  [erro ao ler] {path.name}: {e}")
            return ""


def chunk_text(text: str, size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP):
    """Chunking simples por caracteres com overlap."""
    text = text.strip()
    if not text:
        return []
    chunks = []
    start = 0
    while start < len(text):
        end = start + size
        chunks.append(text[start:end])
        start = end - overlap
    return chunks


def file_hash(path: Path) -> str:
    """Hash do conteúdo do arquivo, pra saber se mudou desde a última ingestão."""
    h = hashlib.sha256()
    h.update(path.read_bytes())
    return h.hexdigest()[:16]


def get_collection():
    client = chromadb.PersistentClient(path=str(VECTOR_DB_PATH))
    return client.get_or_create_collection(name=COLLECTION_NAME)


# Formatos imutáveis: se já foi indexado uma vez, nunca precisa reconferir conteúdo.
IMMUTABLE_EXT = {".pdf"}
# Formatos editáveis: sempre reconferir hash pra detectar mudanças de conteúdo.
EDITABLE_EXT = {".docx", ".odt", ".doc", ".txt", ".md"}


def already_indexed_by_id(collection, doc_id: str) -> bool:
    """Verifica só se o doc_id já existe no índice, sem olhar conteúdo (usado para PDFs)."""
    try:
        existing = collection.get(where={"doc_id": doc_id}, limit=1)
        return bool(existing["ids"])
    except Exception:
        return False


def already_indexed(collection, doc_id: str, current_hash: str) -> bool:
    """Verifica se esse arquivo (com esse hash) já está indexado."""
    try:
        existing = collection.get(where={"doc_id": doc_id}, limit=1)
        if existing["ids"]:
            return existing["metadatas"][0].get("file_hash") == current_hash
    except Exception:
        pass
    return False


def remove_old_chunks(collection, doc_id: str):
    try:
        collection.delete(where={"doc_id": doc_id})
    except Exception:
        pass


def ingest_file(path: Path, root: Path, source_tag: str, collection):
    doc_id = f"{source_tag}::{path.relative_to(root)}"
    suffix = path.suffix.lower()

    # PDF (e outros formatos imutáveis): se já está no índice, pula sem nem calcular hash.
    if suffix in IMMUTABLE_EXT:
        if already_indexed_by_id(collection, doc_id):
            return False
        current_hash = None  # não precisamos de hash real pra formato imutável
    else:
        # Formatos editáveis: sempre reconfere o hash pra detectar mudanças de conteúdo.
        try:
            current_hash = file_hash(path)
        except (FileNotFoundError, OSError, PermissionError) as e:
            print(f"  [pulado — não foi possível ler] {path.relative_to(root)}: {e}")
            return False

        if already_indexed(collection, doc_id, current_hash):
            return False  # nada mudou, pula

    tqdm.write(f"  Indexando [{source_tag}]: {path.relative_to(root)}")
    remove_old_chunks(collection, doc_id)

    # Timeout aqui protege contra PDF corrompido travando o pypdf — não limita arquivos grandes normais.
    try:
        text = run_with_timeout(extract_text, EXTRACTION_TIMEOUT_SECONDS, path)
    except FutureTimeoutError:
        log_skipped(path, f"extração de texto travou (>{EXTRACTION_TIMEOUT_SECONDS}s) — provável PDF corrompido")
        return False
    except (FileNotFoundError, OSError, PermissionError) as e:
        log_skipped(path, f"erro na extração — {e}")
        return False

    chunks = chunk_text(text)
    if not chunks:
        return False

    ids, embeddings, documents, metadatas = [], [], [], []
    chunk_bar = tqdm(
        enumerate(chunks), total=len(chunks),
        desc=f"    {path.name[:40]}", unit="chunk", leave=False
    )
    for i, chunk in chunk_bar:
        # Timeout por CHAMADA individual — um arquivo com muitos chunks pode levar o tempo que precisar,
        # desde que cada chamada ao Ollama continue progredindo dentro do prazo.
        try:
            emb = run_with_timeout(
                lambda c: ollama.embeddings(model=EMBED_MODEL, prompt=c)["embedding"],
                CHUNK_EMBED_TIMEOUT_SECONDS,
                chunk,
            )
        except FutureTimeoutError:
            chunk_bar.close()
            log_skipped(
                path,
                f"embedding travou no chunk {i+1}/{len(chunks)} (>{CHUNK_EMBED_TIMEOUT_SECONDS}s sem resposta do Ollama) — arquivo pulado inteiro"
            )
            return False  # não indexa parcialmente — ou o arquivo inteiro, ou nada

        ids.append(f"{doc_id}::{i}")
        embeddings.append(emb)
        documents.append(chunk)
        metadatas.append({
            "doc_id": doc_id,
            "source_type": source_tag,       # "biblioteca" ou "affine"
            "file_hash": current_hash or "imutavel",
            "source_path": str(path),
            "chunk_index": i,
        })

    collection.add(ids=ids, embeddings=embeddings, documents=documents, metadatas=metadatas)
    return True


def scan_and_ingest(root: Path, source_tag: str):
    if not root.exists():
        print(f"Pasta não encontrada: {root}")
        sys.exit(1)

    collection = get_collection()
    files = [p for p in root.rglob("*") if p.is_file() and p.suffix.lower() in SUPPORTED_EXT]

    if not files:
        print(f"Nenhum arquivo suportado encontrado em {root} (tipos: {SUPPORTED_EXT})")
        return

    print(f"[{source_tag}] Encontrados {len(files)} arquivos. Verificando mudanças...")
    print(f"(timeout de {CHUNK_EMBED_TIMEOUT_SECONDS}s por chamada de embedding — arquivos grandes podem levar o tempo que precisarem, desde que continuem progredindo)")
    updated = 0
    skipped = 0

    file_bar = tqdm(files, desc="Arquivos", unit="arquivo")
    for path in file_bar:
        file_bar.set_postfix_str(path.name[:30])
        try:
            if ingest_file(path, root, source_tag, collection):
                updated += 1
        except Exception as e:
            skipped += 1
            log_skipped(path, f"erro inesperado — {e}")
            continue

    print(f"Concluído. {updated} arquivo(s) atualizado(s)/adicionado(s), {skipped} pulado(s) por erro [{source_tag}].")
    print(f"Total de chunks na coleção (todas as fontes): {collection.count()}")
    if skipped:
        print(f"Detalhes dos pulados salvos em: {SKIPPED_LOG_PATH}")


def watch_mode(root: Path, source_tag: str, interval=60):
    print(f"Modo watch ativo em [{source_tag}] — reingestão a cada {interval}s. Ctrl+C pra parar.")
    while True:
        scan_and_ingest(root, source_tag)
        time.sleep(interval)


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--root", type=str, default=str(YURIAI / "Biblioteca"),
                         help="Pasta raiz a indexar (padrão: Biblioteca)")
    parser.add_argument("--source-tag", type=str, default="biblioteca",
                         help="Rótulo da fonte, ex: biblioteca, affine")
    parser.add_argument("--watch", action="store_true", help="Fica rodando e reingere periodicamente")
    parser.add_argument("--interval", type=int, default=60, help="Intervalo do watch em segundos")
    args = parser.parse_args()

    root_path = Path(args.root)

    if args.watch:
        watch_mode(root_path, args.source_tag, args.interval)
    else:
        scan_and_ingest(root_path, args.source_tag)
